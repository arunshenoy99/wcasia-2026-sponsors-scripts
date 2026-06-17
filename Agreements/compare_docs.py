#!/usr/bin/env python3
"""Compare the 2025 reference document with the generated 2026 template (see README.md)."""

import os
import sys

from docx import Document
from reference_doc import find_reference_doc

GENERATED_TEMPLATE = "WordCamp Asia 2026 Super Admin Sponsorship Agreement Template.docx"

reference_path = find_reference_doc()
if reference_path is None:
    sys.exit("No 2025 reference agreement (.docx) found in this folder — see README.md.")
if not os.path.exists(GENERATED_TEMPLATE):
    sys.exit(
        f"Generated template '{GENERATED_TEMPLATE}' not found. "
        "Run generate_templates.py first."
    )

print("=== ORIGINAL 2025 DOCUMENT ===\n")
doc_orig = Document(str(reference_path))

for i, para in enumerate(doc_orig.paragraphs[6:30]):
    style_name = para.style.name if para.style else "None"
    is_bold = any(run.bold for run in para.runs) if para.runs else False
    text = para.text[:80] if para.text else ""
    
    # Check for numbering
    num_info = ""
    if para._element.pPr is not None and para._element.pPr.numPr is not None:
        numPr = para._element.pPr.numPr
        ilvl = numPr.ilvl
        numId = numPr.numId
        if ilvl is not None:
            num_info = f" [Level: {ilvl.val}]"
        if numId is not None:
            num_info += f" [NumId: {numId.val}]"
    
    # Check which runs are bold
    bold_runs = []
    if para.runs:
        for j, run in enumerate(para.runs):
            if run.bold:
                bold_runs.append(f"run{j}:{run.text[:30]}")
    
    print(f"{i+1:2d}. Style: {style_name:15s} Bold: {str(is_bold):5s}{num_info}")
    if bold_runs:
        print(f"     Bold runs: {bold_runs}")
    if text:
        print(f"     Text: {text}")
    print()

print("\n=== GENERATED 2026 DOCUMENT ===\n")
doc_new = Document(GENERATED_TEMPLATE)

for i, para in enumerate(doc_new.paragraphs[6:30]):
    style_name = para.style.name if para.style else "None"
    is_bold = any(run.bold for run in para.runs) if para.runs else False
    text = para.text[:80] if para.text else ""
    
    # Check for numbering
    num_info = ""
    if para._element.pPr is not None and para._element.pPr.numPr is not None:
        numPr = para._element.pPr.numPr
        ilvl = numPr.ilvl
        numId = numPr.numId
        if ilvl is not None:
            num_info = f" [Level: {ilvl.val}]"
        if numId is not None:
            num_info += f" [NumId: {numId.val}]"
    
    # Check which runs are bold
    bold_runs = []
    if para.runs:
        for j, run in enumerate(para.runs):
            if run.bold:
                bold_runs.append(f"run{j}:{run.text[:30]}")
    
    print(f"{i+1:2d}. Style: {style_name:15s} Bold: {str(is_bold):5s}{num_info}")
    if bold_runs:
        print(f"     Bold runs: {bold_runs}")
    if text:
        print(f"     Text: {text}")
    print()

