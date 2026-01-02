#!/usr/bin/env python3
"""
Generate WordCamp Asia 2026 Sponsor Agreement Templates
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

def number_to_words(n):
    """Convert number to words for amounts"""
    ones = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE", 
            "TEN", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", 
            "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    
    if n == 0:
        return "ZERO"
    if n < 20:
        return ones[n]
    if n < 100:
        return tens[n // 10] + (" " + ones[n % 10] if n % 10 else "")
    if n < 1000:
        return ones[n // 100] + " HUNDRED" + (" " + number_to_words(n % 100) if n % 100 else "")
    if n < 1000000:
        return number_to_words(n // 1000) + " THOUSAND" + (" " + number_to_words(n % 1000) if n % 1000 else "")
    return ""

def format_amount_words(amount):
    """Format amount in words"""
    dollars = int(amount)
    if dollars == 40000:
        return "FORTY THOUSAND"
    elif dollars == 30000:
        return "THIRTY THOUSAND"
    elif dollars == 20000:
        return "TWENTY THOUSAND"
    elif dollars == 15000:
        return "FIFTEEN THOUSAND"
    elif dollars == 10000:
        return "TEN THOUSAND"
    elif dollars == 5000:
        return "FIVE THOUSAND"
    elif dollars == 2500:
        return "TWO THOUSAND FIVE HUNDRED"
    return number_to_words(dollars)

def add_numbered_point(doc, text, level=0, is_bold=False, num_id=5, bold_text=None):
    """Add a numbered list point with proper formatting
    If bold_text is provided, only that portion will be bold (for main items)
    If is_bold is True and bold_text is None, the whole text will be bold (for section headers)
    """
    para = doc.add_paragraph()
    
    if is_bold and bold_text:
        # Only bold the label portion
        run_bold = para.add_run(bold_text)
        run_bold.bold = True
        # Add the rest as normal
        rest_text = text[len(bold_text):]
        if rest_text:
            para.add_run(rest_text)
    elif is_bold:
        # Bold the whole thing (for section headers)
        run = para.add_run(text)
        run.bold = True
    else:
        # No bold
        para.add_run(text)
    
    # Set up numbering
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Set list level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # Set numbering ID (same for all items in the list)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    
    pPr.append(numPr)
    return para

def add_sub_numbered_point(doc, text):
    """Add a sub-numbered list point (level 2, not bold)"""
    return add_numbered_point(doc, text, level=2, is_bold=False)

def create_template(tier_name, tier_data):
    """Create a template for a specific tier"""
    doc = Document()
    
    # Ensure numbering definitions exist by adding a dummy numbered paragraph
    # This will create the numbering part if it doesn't exist
    dummy_para = doc.add_paragraph("", style='List Number')
    # Remove the dummy paragraph but keep the numbering part
    doc._body._body.remove(dummy_para._element)
    
    # Copy numbering definitions from original document to get the exact format
    try:
        orig_doc = Document("1-Elementor WordCamp Asia 2025 Super Admin Sponsorship Agreement.docx")
        if orig_doc.part.numbering_part and doc.part.numbering_part:
            # Get the XML from original
            orig_numbering = orig_doc.part.numbering_part.element
            new_numbering = doc.part.numbering_part.element
            
            # Clear existing content and copy all children
            for child in list(new_numbering):
                new_numbering.remove(child)
            
            # Copy all children from original
            for child in orig_numbering:
                new_numbering.append(child)
    except Exception as e:
        print(f"Warning: Could not copy numbering definitions: {e}")
    
    # Title
    title = doc.add_heading('WordCamp Sponsorship Agreement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('This Sponsorship Agreement is made by and between')
    doc.add_paragraph('WordPress Community Support, a Public Benefit Corporation (the "WPCS"), with registered address:')
    doc.add_paragraph('660 4th Street #119, San Francisco CA 94107 (Tax ID: 81-0896291),')
    doc.add_paragraph(f'And [SPONSOR_NAME] (the "Sponsor").')
    doc.add_paragraph('This Agreement is effective as of * Date [DATE].')
    
    # Sponsored events
    add_numbered_point(doc, 'Sponsored events. With assistance from its local organizers, WPCS hosts WordCamp conferences throughout the world. This Agreement pertains to a WordCamp event hosted in the following location and time period:', level=0, is_bold=True, bold_text='Sponsored events')
    doc.add_paragraph(' Event: WordCamp Asia 2026  Location: Jio World Convention Centre, Mumbai, India')
    doc.add_paragraph(' Time Period: April 9 - 11, 2026')
    
    # Sponsorship Amount
    amount_words = format_amount_words(tier_data['amount'])
    full_text = f'Sponsorship Amount. Within 30 days of this Agreement or before the start of the first Sponsored WordCamp, whichever is sooner, the Sponsor agrees to pay ${tier_data["amount"]:,} ({amount_words} DOLLARS) to WPCS (the "Sponsorship") for the {tier_name} Sponsorship Tier.'
    add_numbered_point(doc, full_text, level=0, is_bold=True, bold_text='Sponsorship Amount')
    
    # Use of Funds
    add_numbered_point(doc, 'Use of Funds. WPCS will use the Sponsorship to cover its costs and the costs of its volunteers and agents in connection with organizing, promoting, and operating the Sponsored WordCamp(s). Any excess remaining after these costs are paid may be used by WPCS for its unrestricted general support.', level=0, is_bold=True, bold_text='Use of Funds')
    
    # Recognition of Sponsor
    add_numbered_point(doc, 'Recognition of Sponsor at the Sponsored WordCamps. In recognition of its support through the Sponsorship, WPCS will provide the following benefits to the Sponsor:', level=0, is_bold=True, bold_text='Recognition of Sponsor at the Sponsored WordCamps')
    
    # Booth Details
    add_numbered_point(doc, 'Booth Details', level=1, is_bold=True)
    if tier_data.get('no_booth'):
        add_sub_numbered_point(doc, 'No booth is provided.')
    else:
        if tier_data.get('booth_size_sqm'):
            add_sub_numbered_point(doc, f'{tier_data["booth_size_sqm"]} square meters for your sponsor booth for 2 days at a premium position in the sponsor hall.')
        else:
            add_sub_numbered_point(doc, '[BOOTH_SIZE_SQM] square meters for your sponsor booth for 2 days at a premium position in the sponsor hall.')
        if tier_data.get('booth_height'):
            add_sub_numbered_point(doc, f'Maximum height is {tier_data["booth_height"]} meters')
        else:
            add_sub_numbered_point(doc, 'Maximum height is [BOOTH_HEIGHT] meters')
    
    # Brand Awareness
    add_numbered_point(doc, 'Brand Awareness', level=1, is_bold=True)
    
    # Logo details
    if tier_name == "Super Admin":
        add_sub_numbered_point(doc, 'Prominent, largest logo and link on the official WordCamp Asia website')
        add_sub_numbered_point(doc, 'Extra-large logo on shared sponsors banner (or any equivalent standing boards/banners)')
        add_sub_numbered_point(doc, 'Logo on back of volunteer / staff T-shirts')
        add_sub_numbered_point(doc, 'Logo on attendee swag bag (extra-large placement)')
    else:
        logo_size = tier_data['logo_size']
        add_sub_numbered_point(doc, f'{logo_size} logo and link on the official WordCamp Asia website')
        add_sub_numbered_point(doc, f'{logo_size} logo on shared sponsors banner (or any equivalent standing boards/banners)')
    
    # Brand page
    if tier_data['social_media_type'] == 'Dedicated':
        add_sub_numbered_point(doc, f'Dedicated thank you posts for your brand on the WordCamp Asia website with your chosen message ({tier_data["brand_page_words"]} words max.). You can add links to your website and social media channels.')
    elif tier_data['social_media_type'] == 'Shared':
        add_sub_numbered_point(doc, f'Shared thank you posts for your brand on the WordCamp Asia website with your chosen message ({tier_data["brand_page_words"]} words max.). You can add links to your website and social media channels.')
    elif tier_data['social_media_type'] == 'Grouped':
        add_sub_numbered_point(doc, f'Share thank you posts for your brand on the WordCamp Asia website with your chosen message ({tier_data["brand_page_words"]} words max.). You can add links to your website and your social media channels.')
    
    add_sub_numbered_point(doc, 'A dedicated page presenting your product/business on the WordCamp Asia website with your chosen message.')
    
    # Mentions
    if tier_data.get('mentioned_opening') or tier_data.get('mentioned_closing'):
        if tier_data.get('mentioned_opening') and tier_data.get('mentioned_closing'):
            add_sub_numbered_point(doc, 'Acknowledgment and thanks in opening and closing remarks')
        elif tier_data.get('mentioned_opening'):
            add_sub_numbered_point(doc, 'Acknowledgment and thanks in opening remarks')
        elif tier_data.get('mentioned_closing'):
            add_sub_numbered_point(doc, 'Acknowledgment and thanks in closing remarks')
    
    # Social media
    if tier_data['social_media_type'] == 'Dedicated':
        add_sub_numbered_point(doc, 'Dedicated pre-event thank-you social media post')
        add_sub_numbered_point(doc, 'Dedicated post-event social media post')
    elif tier_data['social_media_type'] == 'Shared':
        add_sub_numbered_point(doc, 'Shared pre-event thank-you social media post')
        add_sub_numbered_point(doc, 'Shared post-event social media post')
    elif tier_data['social_media_type'] == 'Grouped':
        add_sub_numbered_point(doc, 'Grouped pre-event social thank-you post')
        add_sub_numbered_point(doc, 'Grouped post-event social post')
    
    if tier_data['social_media_type'] == 'Dedicated':
        add_sub_numbered_point(doc, 'Dedicated official announcement of your sponsorship.')
    else:
        add_sub_numbered_point(doc, 'Shared official announcement of your sponsorship.')
    
    # Swag in attendee bag (Super Admin and Admin)
    if tier_data.get('swag_in_bag'):
        add_sub_numbered_point(doc, 'Swag in attendee bag (provided by sponsor)')
    
    # Video adverts (only for Super Admin and Admin)
    if tier_data.get('video_advert'):
        add_sub_numbered_point(doc, f'Video advert ({tier_data["video_advert"]})')
    
    # Professional photoshoot
    if tier_data.get('professional_photoshoot'):
        add_sub_numbered_point(doc, 'Professional photoshoot')
    
    # Option to add sub-brand
    if tier_data.get('sub_brand_option'):
        add_sub_numbered_point(doc, 'Option to add a sub-brand (via add-ons)')
    
    # Raffle announcement
    if tier_data.get('raffle_announcement'):
        add_sub_numbered_point(doc, 'Option for raffle during closing remarks. Take up to [RAFFLE_TIME] seconds to be an on-stage presence for your brand, building brand loyalty by declaring a raffle winner.')
    
    # Tickets
    add_numbered_point(doc, 'Tickets', level=1, is_bold=True)
    if tier_data.get('viewer_tickets'):
        add_sub_numbered_point(doc, f'{tier_data["tickets_word"]} ({tier_data["staff_tickets"]}) regular attendee ticket for the conference.')
    else:
        add_sub_numbered_point(doc, f'{tier_data["tickets_word"]} ({tier_data["staff_tickets"]}) staff tickets to access your booth before opening and after closing of the venue')
    add_sub_numbered_point(doc, f'{tier_data["social_seats_word"]} ({tier_data["social_night_seats"]}) invite{"s" if tier_data["social_night_seats"] > 1 else ""} to the Social event (the pre-event for speakers, organizers, and sponsors).')
    
    # Career Corner
    if tier_data.get('career_corner'):
        add_numbered_point(doc, 'Career Corner', level=1, is_bold=True)
        if tier_data['career_corner'] == 'Dedicated':
            add_sub_numbered_point(doc, 'Option for career corner. A dedicated space in the sponsor hall for sponsors to host a career fair, focusing on hiring talented individuals from the local community. A table with chairs will be provided. You may bring your own standing banner (One banner).')
        else:
            add_sub_numbered_point(doc, 'Career Corner: Shared space')
    
    # Social Corner
    if tier_data.get('social_corner'):
        add_numbered_point(doc, 'Social Corner', level=1, is_bold=True)
        add_sub_numbered_point(doc, 'A dedicated space in the sponsor hall for sponsor meetings to foster networking and business opportunities.')
    
    # Eligibility (for Subscriber and Viewer)
    if tier_data.get('eligibility'):
        add_numbered_point(doc, 'Eligibility Requirements', level=1, is_bold=True)
        for req in tier_data['eligibility']:
            add_sub_numbered_point(doc, req)
    
    # Standard clauses
    doc.add_paragraph()
    add_numbered_point(doc, 'The Sponsor is responsible for providing to WPCS in a timely manner the links referenced above as well as any name or logo artwork for use in the above acknowledgments. The Sponsor agrees, however, that the specific format of the above acknowledgment (e.g., time-length of slide displays and relative size of the Sponsor\'s logo) will be in WPCS\'s discretion.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Notwithstanding anything else in this Agreement, the Sponsor understands and agrees that any acknowledgment by WPCS of the Sponsor is limited to the terms described in Exhibit A. The Sponsor also understands and agrees that WPCS will not endorse the Sponsor or any product or service offered by the Sponsor, and that nothing in this Agreement provides any right to the Sponsor or its representatives to speak at a Sponsored WordCamp or meetup.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Sponsor Conduct. The Sponsor recognizes that, in associating itself with WPCS and the Sponsored WordCamps, the Sponsor expected to support the WordPress project and its principles. Accordingly, the Sponsor agrees to comply with the Sponsor Guidelines attached as Exhibit A in conducting any activities at or in connection with the Sponsored WordCamps.', level=0, is_bold=True, bold_text='Sponsor Conduct')
    
    add_numbered_point(doc, 'Use of WordCamp names. The Sponsor may in its reasonable discretion use the name and logo of each Sponsored WordCamp, and may refer or link to each Sponsored WordCamp, in any press release, website, advertisement, or other public document or announcement, including without limitation in a general list of the Sponsor\'s supported organizations and as otherwise required by law; provided, however, that any such use must be in compliance with the Sponsor Guidelines attached as Exhibit A (including but not limited to the prohibition on the use of WPCS\'s name to imply any endorsement of the Sponsor\'s products or services).', level=0, is_bold=True, bold_text='Use of WordCamp names')
    
    add_numbered_point(doc, 'Any breach by the Sponsor of Section 5 or Section 6 will constitute a material breach of this Agreement, as a result of which WPCS may terminate this Agreement and retain the Sponsorship for its unrestricted use if the Sponsor does not cure such breach to the reasonable satisfaction of WPCS in a reasonably prompt timeframe under the circumstances (and in any event immediately, if such breach occurs during a WordCamp).', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Trademarks. The Sponsor and WPCS hereby grant each other permission to use the other party\'s name, logo, and other trademarks in connection with the activities contemplated above. These permissions are, however, revocable, non-exclusive, and non-transferable, and each party agrees to use the other party\'s logo or trademark only in accordance with any trademark usage guidelines that the other party may provide from time to time. Neither party will hold the other party liable for any incidental or consequential damages arising from that other party\'s use of its trademarks in connection with this Agreement. Except as expressly provided above, any use of the WordPress trademarks is subject to the WordPress  Trademark Policy listed at http://wordpressfoundation.org/trademark-policy.', level=0, is_bold=True, bold_text='Trademarks')
    
    add_numbered_point(doc, 'Relationship of the Parties. This Agreement is not to be construed as creating any agency, partnership, joint venture, or any other form of association, for tax purposes or otherwise, between the parties, and neither party will make any such representation to anyone. Neither party will have any right or authority, express or implied, to assume or create any obligation of any kind, or to make any representation or warranty, on behalf of the other party or to bind the other party in any respect.', level=0, is_bold=True, bold_text='Relationship of the Parties')
    
    add_numbered_point(doc, 'Governing Law. This Agreement will be governed by and construed in accordance with the laws of the State of California, USA, without reference to its conflict of laws provisions.', level=0, is_bold=True, bold_text='Governing Law')
    
    add_numbered_point(doc, 'Severability. If any provision of this Agreement is held to be invalid, void, or otherwise unenforceable, that provision will be enforced to the maximum extent possible so as to effect the intent of the parties, and the remainder of this Agreement will remain in full force and effect.', level=0, is_bold=True, bold_text='Severability')
    
    add_numbered_point(doc, 'Assignment. Neither WPCS nor the Sponsor will have the right to assign this Agreement without the prior written consent of the other party, and any purported assignment without such consent will be void. WPCS may delegate its duties under this Agreement to its volunteers and local WordCamp organizers.', level=0, is_bold=True, bold_text='Assignment')
    
    add_numbered_point(doc, 'Refund and Cancellation Policy. WordCamp Sponsors will not be acknowledged until payment is received in full. Sponsors may request a refund and cancel their sponsorship within 5 business days of payment of the sponsorship invoice. 5 business days after the sponsorship invoice is paid, refunds are no longer available. If a WordCamp is canceled, sponsors will be refunded their sponsorship fees in full.', level=0, is_bold=True, bold_text='Refund and Cancellation Policy')
    
    add_numbered_point(doc, 'Entire Agreement; Amendment. This Agreement (including Exhibit A) constitutes the entire agreement of WPCS and the Sponsor with respect to the subject matter set forth herein, and this Agreement supersedes any prior or contemporaneous oral or written agreements, understandings, or communications or past courses of dealing between the Sponsor and WPCS with respect to that subject matter. This Agreement may not be amended or modified, except in a written amendment signed by duly authorized representatives of both parties.', level=0, is_bold=True, bold_text='Entire Agreement; Amendment')
    
    add_numbered_point(doc, 'Counterparts. This Agreement may be executed in one or more counterparts, each of which will be deemed an original, but all of which together will constitute one and the same agreement.', level=0, is_bold=True, bold_text='Counterparts')
    
    add_numbered_point(doc, 'Sanctions compliance. The WPCS represents and warrants that neither it nor any of its directors, officers, employees or Ultimate Beneficial Owners (UBOs) are (a) subject to any sanctions, restrictions, or designated on any list of prohibited or restricted parties, including but not limited to the lists maintained by the US Department of the Treasury\'s OFAC, United Nations, European Union, UK Treasury or any other relevant sanctions compliance authorities, or (b) located in, organized under the laws of, or ordinarily resident in, a country or territory that is subject to US comprehensive country or other sanctioned country by US, EU / UN, UK or any other country. If there is any uncertainty regarding potential relations to sanctioned entities or jurisdictions, it is obligatory to disclose such information prior to contract execution. Failure to do so, and subsequent discovery of such relations, constitute a material breach of contract and grounds for termination.', level=0, is_bold=True, bold_text='Sanctions compliance')
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph('The parties have executed this Agreement as of date set forth above.')
    doc.add_paragraph('Sponsor\t\t\t\t\t\t\t\t\tWPCS')
    doc.add_paragraph()
    doc.add_paragraph('Representative name:\t\t\t\t\t\t\tRepresentative name:')
    doc.add_paragraph('Designation:\t\t\t\t\t\t\t\tDesignation:')
    
    # Exhibit A
    doc.add_page_break()
    doc.add_heading('Exhibit A: WordPress Community Event Sponsor Guidelines', 1)
    
    add_numbered_point(doc, 'Sponsor may provide:', level=0, is_bold=True, bold_text='Sponsor may provide')
    add_sub_numbered_point(doc, 'The sponsor\'s name and logo')
    add_sub_numbered_point(doc, 'Slogans that are an established part of the sponsor\'s image')
    add_sub_numbered_point(doc, 'The sponsor\'s brands and trade names')
    add_sub_numbered_point(doc, 'Sponsor contact information (such as telephone numbers, email addresses, and URLs)')
    add_sub_numbered_point(doc, 'Factual (value-neutral) displays of actual products')
    add_sub_numbered_point(doc, 'Displays or handout materials (such as brochures) with factual, non-comparative descriptions or listings of products or services')
    add_sub_numbered_point(doc, 'Price information, or other indications of savings or value, if factual and provable Inducements to purchase or use the Sponsor\'s products or services, for example by providing coupons or discount purchase codes (subject to approval)')
    add_sub_numbered_point(doc, 'Calls to action, such as "visit this site for details", "call now for a special offer", "join our league of savings", etc.')
    
    add_numbered_point(doc, 'Sponsors may not provide:', level=0, is_bold=True, bold_text='Sponsors may not provide')
    add_sub_numbered_point(doc, 'Promotional or marketing material containing comparative messages about the Sponsor, its products or services, such as "the first name in WordPress hosting", "the easiest way to launch your site", or "the best e-commerce plugin"')
    add_sub_numbered_point(doc, 'Claims that WordPress, the WordPress Foundation, WordPress Community Support, meetup organizers, WordCamps, or WordCamp organizers endorse or favor a Sponsor or its products or services (such as "certified WordPress training" or "WordCamp\'s favorite plugin")')
    
    add_numbered_point(doc, 'Sponsors agree that the WordPress Community Support, any subsidiary or related entity of the WordPress Community Support, and WordCamp organizers have the right to request and review sponsor materials in advance of an event, to require changes to any materials in advance, and to require that any materials that do not meet the above expectations be taken down or that any practices that do not meet the above expectations be discontinued during a WordCamp or event. The above restrictions also apply to material placed on any self-serve swag tables reserved for sponsor use.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'All sponsors are expected to support the WordPress project and its principles, including:', level=0, is_bold=True, bold_text='All sponsors are expected to support the WordPress project and its principles, including')
    add_sub_numbered_point(doc, 'No discrimination on the basis of economic or social status, race, color, ethnic origin, national origin, creed, religion, political belief, sex, sexual orientation, marital status, age, or disability.')
    add_sub_numbered_point(doc, 'No incitement to violence or promotion of hate')
    add_sub_numbered_point(doc, 'No spammers')
    add_sub_numbered_point(doc, 'No jerks')
    add_sub_numbered_point(doc, 'Respect the WordPress trademark.')
    add_sub_numbered_point(doc, 'Embrace the WordPress license; If distributing WordPress-derivative works (themes, plugins, WP distros), any person or business officially associated with WordCamp should give their users the same freedoms that WordPress itself provides: 100% GPL or compatible, the same guidelines we follow on WordPress.org.')
    add_sub_numbered_point(doc, 'Don\'t promote companies or people that violate the trademark or distribute WordPress derivative works which aren\'t 100% GPL compatible.')
    
    doc.add_paragraph()
    add_numbered_point(doc, 'Sponsorship is in no way connected to the opportunity to speak at an official WordPress event and does not alter the WordPress or WordCamp trademark usage policy found at http://wordpressfoundation.org/. The WordPress Foundation and any subsidiary or related entity of the Foundation reserve the right to modify the above requirements and expectations at any time by providing written notice to the sponsor.', level=0, is_bold=False)
    
    return doc

def get_number_word(n):
    """Get word for number"""
    words = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten"]
    if n <= 10:
        return words[n]
    return str(n)

# Tier configurations
tiers = {
    "Super Admin": {
        "amount": 40000,
        "booth_size_sqm": None,  # Use placeholder
        "booth_height": None,  # Use placeholder
        "logo_size": "Extra-large",
        "brand_page_words": 400,
        "social_media_type": "Dedicated",
        "swag_in_bag": True,
        "sub_brand_option": True,
        "video_advert": "60 sec, played twice a day across two days",
        "professional_photoshoot": True,
        "career_corner": "Dedicated",
        "social_corner": True,
        "raffle_announcement": True,
        "mentioned_opening": True,
        "mentioned_closing": True,
        "staff_tickets": 10,
        "tickets_word": "Ten",
        "social_night_seats": 10,
        "social_seats_word": "Ten"
    },
    "Admin": {
        "amount": 30000,
        "booth_size_sqm": None,
        "booth_height": None,
        "logo_size": "Large",
        "brand_page_words": 300,
        "social_media_type": "Dedicated",
        "swag_in_bag": True,
        "sub_brand_option": True,
        "video_advert": "30 sec, played once a day for one day",
        "professional_photoshoot": True,
        "career_corner": "Dedicated",
        "social_corner": True,
        "raffle_announcement": True,
        "mentioned_opening": True,
        "mentioned_closing": True,
        "staff_tickets": 8,
        "tickets_word": "Eight",
        "social_night_seats": 8,
        "social_seats_word": "Eight"
    },
    "Editor": {
        "amount": 20000,
        "booth_size_sqm": None,
        "booth_height": None,
        "logo_size": "Medium",
        "brand_page_words": 200,
        "social_media_type": "Shared",
        "sub_brand_option": True,
        "career_corner": "Shared",
        "social_corner": True,
        "raffle_announcement": True,
        "mentioned_opening": True,
        "mentioned_closing": False,
        "staff_tickets": 6,
        "tickets_word": "Six",
        "social_night_seats": 6,
        "social_seats_word": "Six"
    },
    "Author": {
        "amount": 15000,
        "booth_size_sqm": None,
        "booth_height": None,
        "logo_size": "Regular",
        "brand_page_words": 150,
        "social_media_type": "Shared",
        "sub_brand_option": True,
        "career_corner": "Shared",
        "social_corner": True,
        "raffle_announcement": True,
        "mentioned_opening": True,
        "mentioned_closing": False,
        "staff_tickets": 4,
        "tickets_word": "Four",
        "social_night_seats": 4,
        "social_seats_word": "Four"
    },
    "Contributor": {
        "amount": 10000,
        "booth_size_sqm": None,
        "booth_height": None,
        "logo_size": "Small",
        "brand_page_words": 100,
        "social_media_type": "Grouped",
        "career_corner": "Shared",
        "social_corner": True,
        "raffle_announcement": False,
        "mentioned_opening": False,
        "mentioned_closing": False,
        "staff_tickets": 3,
        "tickets_word": "Three",
        "social_night_seats": 3,
        "social_seats_word": "Three"
    },
    "Subscriber": {
        "amount": 5000,
        "booth_size_sqm": None,
        "booth_height": None,
        "logo_size": "Tiny",
        "brand_page_words": 50,
        "social_media_type": "Grouped",
        "career_corner": "Shared",
        "social_corner": True,
        "raffle_announcement": False,
        "mentioned_opening": False,
        "mentioned_closing": False,
        "staff_tickets": 2,
        "tickets_word": "Two",
        "social_night_seats": 2,
        "social_seats_word": "Two",
        "eligibility": [
            "Your company must generate revenue mainly from the WordPress ecosystem or similar open-source technologies",
            "Your revenue estimates for FY 2025 should be projected at less than USD 1 million",
            "You employ fewer than 50 people"
        ]
    },
    "Viewer": {
        "amount": 2500,
        "no_booth": True,
        "logo_size": "Nano",
        "brand_page_words": 25,
        "social_media_type": "Grouped",
        "career_corner": None,
        "social_corner": False,
        "raffle_announcement": False,
        "mentioned_opening": False,
        "mentioned_closing": False,
        "staff_tickets": 1,
        "tickets_word": "One",
        "viewer_tickets": True,
        "social_night_seats": 1,
        "social_seats_word": "One",
        "eligibility": [
            "Your company must generate revenue mainly from the WordPress ecosystem or similar open-source technologies",
            "Your revenue estimates for FY 2025 should be projected at less than USD 100 thousand"
        ]
    }
}

def create_addon_template():
    """Create a template for addon agreements"""
    doc = Document()
    
    # Ensure numbering definitions exist by adding a dummy numbered paragraph
    dummy_para = doc.add_paragraph("", style='List Number')
    doc._body._body.remove(dummy_para._element)
    
    # Copy numbering definitions from original document
    try:
        orig_doc = Document("1-Elementor WordCamp Asia 2025 Super Admin Sponsorship Agreement.docx")
        if orig_doc.part.numbering_part and doc.part.numbering_part:
            orig_numbering = orig_doc.part.numbering_part.element
            new_numbering = doc.part.numbering_part.element
            
            for child in list(new_numbering):
                new_numbering.remove(child)
            
            for child in orig_numbering:
                new_numbering.append(child)
    except Exception as e:
        print(f"Warning: Could not copy numbering definitions: {e}")
    
    # Title
    title = doc.add_heading('WordCamp Sponsorship Agreement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('This Sponsorship Agreement is made by and between')
    doc.add_paragraph('WordPress Community Support, a Public Benefit Corporation (the "WPCS"), with registered address:')
    doc.add_paragraph('660 4th Street #119, San Francisco CA 94107 (Tax ID: 81-0896291),')
    doc.add_paragraph('And [SPONSOR_NAME] (the "Sponsors").')
    doc.add_paragraph('This Agreement is effective as of * Date [DATE].')
    
    # Sponsored events
    add_numbered_point(doc, 'Sponsored events. With assistance from its local organizers, WPCS hosts WordCamp conferences throughout the world. This Agreement pertains to a WordCamp event hosted in the following location and time period:', level=0, is_bold=True, bold_text='Sponsored events')
    doc.add_paragraph(' Event: WordCamp Asia 2026  Location: Jio World Convention Centre, Mumbai, India')
    doc.add_paragraph(' Time Period: April 9 - 11, 2026')
    
    # Sponsorship Amount
    add_numbered_point(doc, 'Sponsorship Amount. Within 30 days of this Agreement or before the start of the first Sponsored WordCamp, whichever is sooner, the Sponsor agrees to pay the following sponsorships to WPCS:', level=0, is_bold=True, bold_text='Sponsorship Amount')
    doc.add_paragraph('$[AMOUNT_1] ([AMOUNT_1_WORDS] DOLLARS) (the "Sponsorship") for the [ADDON_ITEM_1]')
    doc.add_paragraph('$[AMOUNT_2] ([AMOUNT_2_WORDS] DOLLARS) (the "Sponsorship") for the [ADDON_ITEM_2].')
    
    # Use of Funds
    add_numbered_point(doc, 'Use of Funds. WPCS will use the Sponsorship to cover its costs and the costs of its volunteers and agents in connection with organizing, promoting, and operating the Sponsored WordCamp(s). Any excess remaining after these costs are paid may be used by WPCS for its unrestricted general support.', level=0, is_bold=True, bold_text='Use of Funds')
    
    # Recognition of Sponsor
    add_numbered_point(doc, 'Recognition of Sponsor at the Sponsored WordCamps. In recognition of its support through the Sponsorship, WPCS will provide the following benefits to the Sponsor:', level=0, is_bold=True, bold_text='Recognition of Sponsor at the Sponsored WordCamps')
    doc.add_paragraph()  # Empty line for benefits to be added
    
    # Standard clauses
    doc.add_paragraph()
    add_numbered_point(doc, 'The Sponsor is responsible for providing to WPCS in a timely manner the links referenced above as well as any name or logo artwork for use in the above acknowledgments. The Sponsor agrees, however, that the specific format of the above acknowledgment (e.g., time-length of slide displays and relative size of the Sponsor\'s logo) will be in WPCS\'s discretion.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Notwithstanding anything else in this Agreement, the Sponsor understands and agrees that any acknowledgment by WPCS of the Sponsor is limited to the terms described in Exhibit A. The Sponsor also understands and agrees that WPCS will not endorse the Sponsor or any product or service offered by the Sponsor, and that nothing in this Agreement provides any right to the Sponsor or its representatives to speak at a Sponsored WordCamp or meetup.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Sponsor Conduct. The Sponsor recognizes that, in associating itself with WPCS and the Sponsored WordCamps, the Sponsor expected to support the WordPress project and its principles. Accordingly, the Sponsor agrees to comply with the Sponsor Guidelines attached as Exhibit A in conducting any activities at or in connection with the Sponsored WordCamps.', level=0, is_bold=True, bold_text='Sponsor Conduct')
    
    add_numbered_point(doc, 'Use of WordCamp names. The Sponsor may in its reasonable discretion use the name and logo of each Sponsored WordCamp, and may refer or link to each Sponsored WordCamp, in any press release, website, advertisement, or other public document or announcement, including without limitation in a general list of the Sponsor\'s supported organizations and as otherwise required by law; provided, however, that any such use must be in compliance with the Sponsor Guidelines attached as Exhibit A (including but not limited to the prohibition on the use of WPCS\'s name to imply any endorsement of the Sponsor\'s products or services).', level=0, is_bold=True, bold_text='Use of WordCamp names')
    
    add_numbered_point(doc, 'Any breach by the Sponsor of Section 5 or Section 6 will constitute a material breach of this Agreement, as a result of which WPCS may terminate this Agreement and retain the Sponsorship for its unrestricted use if the Sponsor does not cure such breach to the reasonable satisfaction of WPCS in a reasonably prompt timeframe under the circumstances (and in any event immediately, if such breach occurs during a WordCamp).', level=0, is_bold=False)
    
    add_numbered_point(doc, 'Trademarks. The Sponsor and WPCS hereby grant each other permission to use the other party\'s name, logo, and other trademarks in connection with the activities contemplated above. These permissions are, however, revocable, non-exclusive, and non-transferable, and each party agrees to use the other party\'s logo or trademark only in accordance with any trademark usage guidelines that the other party may provide from time to time. Neither party will hold the other party liable for any incidental or consequential damages arising from that other party\'s use of its trademarks in connection with this Agreement. Except as expressly provided above, any use of the WordPress trademarks is subject to the WordPress  Trademark Policy listed at http://wordpressfoundation.org/trademark-policy.', level=0, is_bold=True, bold_text='Trademarks')
    
    add_numbered_point(doc, 'Relationship of the Parties. This Agreement is not to be construed as creating any agency, partnership, joint venture, or any other form of association, for tax purposes or otherwise, between the parties, and neither party will make any such representation to anyone. Neither party will have any right or authority, express or implied, to assume or create any obligation of any kind, or to make any representation or warranty, on behalf of the other party or to bind the other party in any respect.', level=0, is_bold=True, bold_text='Relationship of the Parties')
    
    add_numbered_point(doc, 'Governing Law. This Agreement will be governed by and construed in accordance with the laws of the State of California, USA, without reference to its conflict of laws provisions.', level=0, is_bold=True, bold_text='Governing Law')
    
    add_numbered_point(doc, 'Severability. If any provision of this Agreement is held to be invalid, void, or otherwise unenforceable, that provision will be enforced to the maximum extent possible so as to effect the intent of the parties, and the remainder of this Agreement will remain in full force and effect.', level=0, is_bold=True, bold_text='Severability')
    
    add_numbered_point(doc, 'Assignment. Neither WPCS nor the Sponsor will have the right to assign this Agreement without the prior written consent of the other party, and any purported assignment without such consent will be void. WPCS may delegate its duties under this Agreement to its volunteers and local WordCamp organizers.', level=0, is_bold=True, bold_text='Assignment')
    
    add_numbered_point(doc, 'Refund and Cancellation Policy. WordCamp Sponsors will not be acknowledged until payment is received in full. Sponsors may request a refund and cancel their sponsorship within 5 business days of payment of the sponsorship invoice. 5 business days after the sponsorship invoice is paid, refunds are no longer available. If a WordCamp is canceled, sponsors will be refunded their sponsorship fees in full.', level=0, is_bold=True, bold_text='Refund and Cancellation Policy')
    
    add_numbered_point(doc, 'Entire Agreement; Amendment. This Agreement (including Exhibit A) constitutes the entire agreement of WPCS and the Sponsor with respect to the subject matter set forth herein, and this Agreement supersedes any prior or contemporaneous oral or written agreements, understandings, or communications or past courses of dealing between the Sponsor and WPCS with respect to that subject matter. This Agreement may not be amended or modified, except in a written amendment signed by duly authorized representatives of both parties.', level=0, is_bold=True, bold_text='Entire Agreement; Amendment')
    
    add_numbered_point(doc, 'Counterparts. This Agreement may be executed in one or more counterparts, each of which will be deemed an original, but all of which together will constitute one and the same agreement.', level=0, is_bold=True, bold_text='Counterparts')
    
    add_numbered_point(doc, 'Sanctions compliance. The WPCS represents and warrants that neither it nor any of its directors, officers, employees or Ultimate Beneficial Owners (UBOs) are (a) subject to any sanctions, restrictions, or designated on any list of prohibited or restricted parties, including but not limited to the lists maintained by the US Department of the Treasury\'s OFAC, United Nations, European Union, UK Treasury or any other relevant sanctions compliance authorities, or (b) located in, organized under the laws of, or ordinarily resident in, a country or territory that is subject to US comprehensive country or other sanctioned country by US, EU / UN, UK or any other country. If there is any uncertainty regarding potential relations to sanctioned entities or jurisdictions, it is obligatory to disclose such information prior to contract execution. Failure to do so, and subsequent discovery of such relations, constitute a material breach of contract and grounds for termination.', level=0, is_bold=True, bold_text='Sanctions compliance')
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph('The parties have executed this Agreement as of date set forth above.')
    doc.add_paragraph('Sponsor\t\t\t\t\t\t\t\t\tWPCS')
    doc.add_paragraph()
    doc.add_paragraph('Representative name:\t\t\t\t\t\t\tRepresentative name:')
    doc.add_paragraph('Designation:\t\t\t\t\t\t\t\tDesignation:')
    
    # Exhibit A
    doc.add_page_break()
    doc.add_heading('Exhibit A: WordPress Community Event Sponsor Guidelines', 1)
    
    add_numbered_point(doc, 'Sponsor may provide:', level=0, is_bold=True, bold_text='Sponsor may provide')
    add_sub_numbered_point(doc, 'The sponsor\'s name and logo')
    add_sub_numbered_point(doc, 'Slogans that are an established part of the sponsor\'s image')
    add_sub_numbered_point(doc, 'The sponsor\'s brands and trade names')
    add_sub_numbered_point(doc, 'Sponsor contact information (such as telephone numbers, email addresses, and URLs)')
    add_sub_numbered_point(doc, 'Factual (value-neutral) displays of actual products')
    add_sub_numbered_point(doc, 'Displays or handout materials (such as brochures) with factual, non-comparative descriptions or listings of products or services')
    add_sub_numbered_point(doc, 'Price information, or other indications of savings or value, if factual and provable Inducements to purchase or use the Sponsor\'s products or services, for example by providing coupons or discount purchase codes (subject to approval)')
    add_sub_numbered_point(doc, 'Calls to action, such as "visit this site for details", "call now for a special offer", "join our league of savings", etc.')
    
    add_numbered_point(doc, 'Sponsors may not provide:', level=0, is_bold=True, bold_text='Sponsors may not provide')
    add_sub_numbered_point(doc, 'Promotional or marketing material containing comparative messages about the Sponsor, its products or services, such as "the first name in WordPress hosting", "the easiest way to launch your site", or "the best e-commerce plugin"')
    add_sub_numbered_point(doc, 'Claims that WordPress, the WordPress Foundation, WordPress Community Support, meetup organizers, WordCamps, or WordCamp organizers endorse or favor a Sponsor or its products or services (such as "certified WordPress training" or "WordCamp\'s favorite plugin")')
    
    add_numbered_point(doc, 'Sponsors agree that the WordPress Community Support, any subsidiary or related entity of the WordPress Community Support, and WordCamp organizers have the right to request and review sponsor materials in advance of an event, to require changes to any materials in advance, and to require that any materials that do not meet the above expectations be taken down or that any practices that do not meet the above expectations be discontinued during a WordCamp or event. The above restrictions also apply to material placed on any self-serve swag tables reserved for sponsor use.', level=0, is_bold=False)
    
    add_numbered_point(doc, 'All sponsors are expected to support the WordPress project and its principles, including:', level=0, is_bold=True, bold_text='All sponsors are expected to support the WordPress project and its principles, including')
    add_sub_numbered_point(doc, 'No discrimination on the basis of economic or social status, race, color, ethnic origin, national origin, creed, religion, political belief, sex, sexual orientation, marital status, age, or disability.')
    add_sub_numbered_point(doc, 'No incitement to violence or promotion of hate')
    add_sub_numbered_point(doc, 'No spammers')
    add_sub_numbered_point(doc, 'No jerks')
    add_sub_numbered_point(doc, 'Respect the WordPress trademark.')
    add_sub_numbered_point(doc, 'Embrace the WordPress license; If distributing WordPress-derivative works (themes, plugins, WP distros), any person or business officially associated with WordCamp should give their users the same freedoms that WordPress itself provides: 100% GPL or compatible, the same guidelines we follow on WordPress.org.')
    add_sub_numbered_point(doc, 'Don\'t promote companies or people that violate the trademark or distribute WordPress derivative works which aren\'t 100% GPL compatible.')
    
    doc.add_paragraph()
    add_numbered_point(doc, 'Sponsorship is in no way connected to the opportunity to speak at an official WordPress event and does not alter the WordPress or WordCamp trademark usage policy found at http://wordpressfoundation.org/. The WordPress Foundation and any subsidiary or related entity of the Foundation reserve the right to modify the above requirements and expectations at any time by providing written notice to the sponsor.', level=0, is_bold=False)
    
    return doc

if __name__ == "__main__":
    for tier_name, tier_data in tiers.items():
        print(f"Generating {tier_name} template...")
        doc = create_template(tier_name, tier_data)
        filename = f"WordCamp Asia 2026 {tier_name} Sponsorship Agreement Template.docx"
        doc.save(filename)
        print(f"Created: {filename}")
    
    # Generate addon agreement template
    print("Generating Addon Agreement template...")
    addon_doc = create_addon_template()
    addon_filename = "WordCamp Asia 2026 Addon Agreement Template.docx"
    addon_doc.save(addon_filename)
    print(f"Created: {addon_filename}")

